#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del Informe Técnico Profesional — Inventario Agrícola Multibodega
Produce: .docx editable + .pdf
"""

import os
import textwrap
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak,
    Table, TableStyle, KeepTogether
)

BASE = Path(__file__).resolve().parent
ASSETS = BASE / "assets_informe"
OUT = BASE / "entregables"
ASSETS.mkdir(parents=True, exist_ok=True)
OUT.mkdir(parents=True, exist_ok=True)

AUTHOR = "José Cabrera"
INSTITUTION = "Proyecto Inventario Agrícola"
DATE = datetime.now().strftime("%d de %B de %Y")
PROJECT = "Inventario Agrícola Multibodega"
SUBTITLE = "Control de Inventario • Firebase • Room • Offline-First"

# ─── Diagram generation ─────────────────────────────────────────────────────

def _save_fig(name, fig):
    path = ASSETS / name
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def diagram_mvvm():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    boxes = [
        (1, 4.5, "UI Layer\n(Jetpack Compose)", "#2E7D32"),
        (1, 3.0, "ViewModel\n(StateFlow / MVVM)", "#1565C0"),
        (1, 1.5, "Repository / Service\n(MovimientoInventarioService)", "#6A1B9A"),
        (6, 1.5, "Room Database\n(SQLite local)", "#EF6C00"),
        (6, 3.0, "Firebase RTDB\n(Nube)", "#C62828"),
    ]
    for x, y, t, c in boxes:
        ax.add_patch(FancyBboxPatch((x, y), 3.2, 0.9, boxstyle="round,pad=0.05", fc=c, ec="white", alpha=0.9))
        ax.text(x + 1.6, y + 0.45, t, ha="center", va="center", color="white", fontsize=9, fontweight="bold")
    for y1, y2 in [(4.5, 3.9), (3.0, 2.4), (1.5, 1.5)]:
        ax.annotate("", xy=(2.6, y2), xytext=(2.6, y1), arrowprops=dict(arrowstyle="->", color="#333", lw=2))
    ax.annotate("", xy=(4.2, 1.95), xytext=(6, 1.95), arrowprops=dict(arrowstyle="<->", color="#333", lw=2))
    ax.annotate("", xy=(7.6, 2.4), xytext=(7.6, 3.0), arrowprops=dict(arrowstyle="<->", color="#333", lw=2))
    ax.text(5, 5.5, "Arquitectura MVVM — Inventario Agrícola", ha="center", fontsize=14, fontweight="bold")
    return _save_fig("01_arquitectura_mvvm.png", fig)


def diagram_offline_first():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    steps = ["Usuario\noperación", "ViewModel", "Room\n(escritura\ninmediata)", "¿Online?", "Firebase\npush", "SyncQueue\n(reintentos)", "Network\nCallback\npull"]
    xs = [0.3, 1.8, 3.5, 5.2, 6.8, 8.2, 9.5]
    for i, (x, s) in enumerate(zip(xs, steps)):
        c = "#2E7D32" if i in (2, 4) else "#1565C0"
        ax.add_patch(FancyBboxPatch((x - 0.55, 2), 1.1, 1.2, boxstyle="round", fc=c, ec="white", alpha=0.85))
        ax.text(x, 2.6, s, ha="center", va="center", color="white", fontsize=7, fontweight="bold")
        if i < len(xs) - 1:
            ax.annotate("", xy=(xs[i+1]-0.55, 2.6), xytext=(x+0.55, 2.6), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.text(5, 4.3, "Flujo Offline-First con sincronización bidireccional", ha="center", fontsize=13, fontweight="bold")
    return _save_fig("02_offline_first.png", fig)


def diagram_erd():
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis("off")
    entities = {
        "Bodega": (0.5, 6, ["id PK", "codigoCorto", "nombre"]),
        "Producto": (3.5, 6, ["id PK", "codigo", "bodegaId FK", "cantidad", "stockMinimo"]),
        "Categoria": (6.5, 6, ["id PK", "prefijo", "correlativoActual"]),
        "Entrada": (0.5, 3.5, ["id PK", "codigoProducto FK", "bodegaId FK", "cantidad"]),
        "Salida": (3.5, 3.5, ["id PK", "codigoProducto FK", "bodegaId FK", "cantidad"]),
        "Kardex": (6.5, 3.5, ["id PK", "codigoProducto", "tipoMovimiento", "saldoNuevo"]),
        "Factura": (9.5, 3.5, ["id PK", "numeroFactura", "bodegaId FK"]),
        "Usuario": (0.5, 1, ["id PK", "uuid", "rol", "username"]),
        "Auditoria": (3.5, 1, ["id PK", "productoId FK", "stockFisico", "estado"]),
        "Vale": (6.5, 1, ["idVale PK", "codigoVale", "bodegaId FK"]),
    }
    for name, (x, y, fields) in entities.items():
        h = 0.35 + len(fields) * 0.28
        ax.add_patch(FancyBboxPatch((x, y), 2.4, h, boxstyle="square,pad=0.02", fc="#E8F5E9", ec="#2E7D32", lw=1.5))
        ax.text(x + 1.2, y + h - 0.2, name, ha="center", fontweight="bold", fontsize=9)
        for i, f in enumerate(fields):
            ax.text(x + 0.1, y + h - 0.55 - i*0.28, f, fontsize=7)
    # crow's foot style lines
    ax.plot([2.9, 3.5], [6.8, 6.8], "k-", lw=1)
    ax.plot([2.9, 3.5], [3.9, 3.9], "k-", lw=1)
    ax.text(3.1, 6.95, "1:N", fontsize=8)
    ax.text(5, 8.2, "Diagrama Entidad-Relación (ERD) — Room Database v30", ha="center", fontsize=13, fontweight="bold")
    return _save_fig("03_erd_database.png", fig)


def diagram_entrada_salida():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(2.5, 5.3, "FLUJO ENTRADA", ha="center", fontweight="bold", color="#2E7D32", fontsize=12)
    ax.text(7.5, 5.3, "FLUJO SALIDA", ha="center", fontweight="bold", color="#C62828", fontsize=12)
    ent = ["Formulario\nCrearEntradas", "EntradaViewModel", "MovimientoInventario\nService", "↑ Stock\n↑ Kardex\n↑ Factura", "Firebase Sync"]
    sal = ["Formulario\nCrearSalidas", "SalidaViewModel", "MovimientoInventario\nService", "↓ Stock\n↓ Kardex", "Alerta\nStock Bajo"]
    for i, (e, s) in enumerate(zip(ent, sal)):
        ye = 4.5 - i * 0.95
        ax.add_patch(FancyBboxPatch((0.8, ye), 3.4, 0.7, boxstyle="round", fc="#2E7D32", alpha=0.75+0.05*i))
        ax.text(2.5, ye+0.35, e, ha="center", va="center", color="white", fontsize=8, fontweight="bold")
        ax.add_patch(FancyBboxPatch((5.8, ye), 3.4, 0.7, boxstyle="round", fc="#C62828", alpha=0.75+0.05*i))
        ax.text(7.5, ye+0.35, s, ha="center", va="center", color="white", fontsize=8, fontweight="bold")
        if i < 4:
            ax.annotate("", xy=(2.5, ye-0.05), xytext=(2.5, ye), arrowprops=dict(arrowstyle="->", color="#333"))
            ax.annotate("", xy=(7.5, ye-0.05), xytext=(7.5, ye), arrowprops=dict(arrowstyle="->", color="#333"))
    return _save_fig("04_flujo_entradas_salidas.png", fig)


def diagram_firebase_sync():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(5, 5.5, "Sincronización Firebase Realtime Database", ha="center", fontsize=13, fontweight="bold")
    tree = "/bodegas/{codigo}/{bodegaId}/\n  productos, entradas, salidas, kardex,\n  facturas, vales, auditorias, traslados"
    ax.add_patch(FancyBboxPatch((0.5, 3.2), 4, 1.8, boxstyle="round", fc="#FFEBEE", ec="#C62828"))
    ax.text(2.5, 4.1, "Firebase RTDB\n" + tree, ha="center", va="center", fontsize=8)
    ax.add_patch(FancyBboxPatch((5.5, 3.2), 4, 1.8, boxstyle="round", fc="#E3F2FD", ec="#1565C0"))
    ax.text(7.5, 4.1, "Room SQLite\ninventario_db v30\n17 DAOs", ha="center", va="center", fontsize=9)
    ax.annotate("Push\n(guardar*)", xy=(5.5, 4.1), xytext=(4.5, 4.1), arrowprops=dict(arrowstyle="<->", lw=2))
    ax.annotate("Pull\n(ValueEventListener\nCloudSyncManager)", xy=(5.5, 3.5), xytext=(4.5, 3.5), arrowprops=dict(arrowstyle="<->", lw=2))
    ax.add_patch(FancyBboxPatch((2, 1), 6, 1.2, boxstyle="round", fc="#FFF3E0", ec="#EF6C00"))
    ax.text(5, 1.6, "SyncQueueManager + OfflineSyncObserver + ConflictResolver", ha="center", fontsize=9, fontweight="bold")
    return _save_fig("05_firebase_sync.png", fig)


def diagram_roles():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    roles = [("ADMIN", "#2E7D32", "Operación completa\nUsuarios, papelera, sync"), 
             ("AUDITOR", "#1565C0", "Supervisión total\nAuditoría global"),
             ("VISOR", "#757575", "Solo lectura\nPanel, kardex, export")]
    for i, (r, c, d) in enumerate(roles):
        x = 0.8 + i * 3.1
        ax.add_patch(FancyBboxPatch((x, 2), 2.6, 2, boxstyle="round", fc=c, alpha=0.85))
        ax.text(x+1.3, 3.5, r, ha="center", color="white", fontsize=12, fontweight="bold")
        ax.text(x+1.3, 2.7, d, ha="center", color="white", fontsize=8)
    ax.text(5, 4.5, "Roles y Permisos — RoleManager + SessionManager", ha="center", fontsize=13, fontweight="bold")
    return _save_fig("06_roles_permisos.png", fig)


def diagram_codigos_automaticos():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    ax.text(5, 3.5, "Sistema de Códigos Automáticos — Prefijo + Correlativo", ha="center", fontsize=13, fontweight="bold")
    flow = ["Categoría\n(Tornillos)", "Prefijo\nTOR", "incrementarCorrelativo()\nRoom + Firebase", "Código\nTOR-0001"]
    for i, f in enumerate(flow):
        x = 0.5 + i * 2.3
        ax.add_patch(FancyBboxPatch((x, 1.2), 2, 1.5, boxstyle="round", fc="#2E7D32", alpha=0.8))
        ax.text(x+1, 1.95, f, ha="center", va="center", color="white", fontsize=9, fontweight="bold")
        if i < 3:
            ax.annotate("", xy=(x+2.05, 1.95), xytext=(x+2, 1.95), arrowprops=dict(arrowstyle="->", lw=2))
    return _save_fig("07_codigos_automaticos.png", fig)


def diagram_navgraph():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(5, 5.5, "Navegación — NavGraph.kt", ha="center", fontsize=13, fontweight="bold")
    nodes = ["Splash", "Login", "MenuP", "MenuBodega", "Inventario", "Entradas", "Salidas", "Dashboard", "Auditoría", "Config"]
    positions = [(1,4),(3,4),(5,4),(7,4),(1,2),(3,2),(5,2),(7,2),(1,0.5),(5,0.5)]
    for (n, (x,y)) in zip(nodes, positions):
        ax.add_patch(FancyBboxPatch((x-0.6, y-0.3), 1.2, 0.6, boxstyle="round", fc="#1565C0", alpha=0.8))
        ax.text(x, y, n, ha="center", va="center", color="white", fontsize=7, fontweight="bold")
    ax.annotate("", xy=(5,4), xytext=(3.6,4), arrowprops=dict(arrowstyle="->", lw=1))
    ax.annotate("", xy=(7,4), xytext=(5.6,4), arrowprops=dict(arrowstyle="->", lw=1))
    return _save_fig("08_navgraph.png", fig)


def diagram_stock_bajo():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    ax.text(5, 3.5, "Flujo Alerta Stock Bajo", ha="center", fontsize=13, fontweight="bold")
    steps = ["Salida\nregistrada", "Stock\nactualizado", "cantidad\n<= stockMinimo", "NotificationHelper\nSTOCK_BAJO", "Canal\nstock_channel\n+ sonido alerta"]
    for i, s in enumerate(steps):
        x = 0.4 + i * 1.85
        c = "#C62828" if i >= 3 else "#EF6C00"
        ax.add_patch(FancyBboxPatch((x, 1.2), 1.6, 1.3, boxstyle="round", fc=c, alpha=0.85))
        ax.text(x+0.8, 1.85, s, ha="center", va="center", color="white", fontsize=7, fontweight="bold")
        if i < 4:
            ax.annotate("", xy=(x+1.65, 1.85), xytext=(x+1.6, 1.85), arrowprops=dict(arrowstyle="->", lw=1.5))
    return _save_fig("09_stock_bajo.png", fig)


def diagram_export():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    ax.text(5, 3.5, "Exportación PDF y Excel", ha="center", fontsize=13, fontweight="bold")
    ax.add_patch(FancyBboxPatch((0.5, 1), 3, 1.8, boxstyle="round", fc="#2E7D32", alpha=0.85))
    ax.text(2, 1.9, "Apache POI 5.2.5\nXSSFWorkbook\nExcel .xlsx", ha="center", color="white", fontsize=9, fontweight="bold")
    ax.add_patch(FancyBboxPatch((3.5, 1), 3, 1.8, boxstyle="round", fc="#1565C0", alpha=0.85))
    ax.text(5, 1.9, "PdfDocument\n(Canvas API)\n+ BrandingExports", ha="center", color="white", fontsize=9, fontweight="bold")
    ax.add_patch(FancyBboxPatch((6.5, 1), 3, 1.8, boxstyle="round", fc="#6A1B9A", alpha=0.85))
    ax.text(8, 1.9, "FileProvider\nCompartir\nWhatsApp / Drive", ha="center", color="white", fontsize=9, fontweight="bold")
    return _save_fig("10_exportacion.png", fig)


def mock_ui_screen(title, items, color, fname):
    fig, ax = plt.subplots(figsize=(5, 8))
    ax.set_xlim(0, 5); ax.set_ylim(0, 8); ax.axis("off")
    ax.add_patch(FancyBboxPatch((0.2, 0.2), 4.6, 7.6, boxstyle="round", fc="#F5F5F5", ec="#333", lw=2))
    ax.add_patch(FancyBboxPatch((0.2, 7), 4.6, 0.8, boxstyle="round", fc=color, ec=color))
    ax.text(2.5, 7.4, title, ha="center", va="center", color="white", fontsize=11, fontweight="bold")
    for i, item in enumerate(items):
        y = 6.2 - i * 0.65
        ax.add_patch(FancyBboxPatch((0.5, y), 4, 0.5, boxstyle="round", fc="white", ec="#CCC"))
        ax.text(0.7, y+0.25, item, fontsize=8, va="center")
    return _save_fig(fname, fig)


def generate_all_diagrams():
    paths = [
        diagram_mvvm(), diagram_offline_first(), diagram_erd(),
        diagram_entrada_salida(), diagram_firebase_sync(), diagram_roles(),
        diagram_codigos_automaticos(), diagram_navgraph(), diagram_stock_bajo(), diagram_export(),
        mock_ui_screen("Login Inventario Agrícola", ["Usuario / Correo", "Contraseña", "Recordar sesión", "Ingresar → Menu Principal"], "#2E7D32", "ui_login.png"),
        mock_ui_screen("Dashboard KPIs", ["Productos activos", "Entradas del mes", "Salidas del mes", "Gráfica área / donut", "Alertas stock bajo"], "#1565C0", "ui_dashboard.png"),
        mock_ui_screen("Inventario", ["Filtros categoría/estado", "Lista productos consulta", "Export PDF / Excel", "Kardex · Stock bajo"], "#2E7D32", "ui_inventario.png"),
        mock_ui_screen("Entradas", ["Nueva entrada / producto", "Factura · Proveedor", "Código auto TOR-0001", "Kardex automático"], "#388E3C", "ui_entradas.png"),
        mock_ui_screen("Salidas", ["Consumo operativo", "Vale / destino", "Validación stock", "Alertas sonoras"], "#C62828", "ui_salidas.png"),
    ]
    return paths


# ─── Word document ────────────────────────────────────────────────────────────

def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(6)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        hs.font.bold = True
        hs.font.size = Pt(16 - level)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)
    r2 = p.add_run(f"  |  {PROJECT}  |  Informe Técnico")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(10)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2); run._r.append(fldChar3)
    note = doc.add_paragraph("(Actualice el índice en Word: clic derecho → Actualizar campo)")
    note.runs[0].italic = True; note.runs[0].font.size = Pt(10)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    for block in text.strip().split("\n\n"):
        doc.add_paragraph(block.strip())


def add_image(doc, path, width=Inches(5.8)):
    if Path(path).exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
            for p in t.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"; r.font.size = Pt(11)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3); section.right_margin = Cm(2.5)
    add_page_number_footer(section)

    # PORTADA
    for _ in range(6): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(PROJECT); r.bold = True; r.font.size = Pt(22); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = st.add_run(SUBTITLE); rs.font.size = Pt(14); rs.font.name = "Times New Roman"
    doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"\n\nDocumentación Técnica Integral\n\nAutor: {AUTHOR}\n{INSTITUTION}\n{DATE}\n\nVersión 1.0")
    rm.font.name = "Times New Roman"; rm.font.size = Pt(12)
    doc.add_page_break()

    # ÍNDICE
    add_heading(doc, "Índice", 1)
    add_toc(doc)
    doc.add_page_break()

    # 1 RESUMEN EJECUTIVO
    add_heading(doc, "1. Resumen Ejecutivo", 1)
    add_para(doc, """
El presente informe documenta de manera integral el sistema Inventario Agrícola Multibodega, una aplicación móvil empresarial desarrollada en Android Studio con Kotlin. El software permite administrar inventarios agrícolas distribuidos en múltiples bodegas, registrar entradas y salidas de materiales, controlar stock mínimo, generar kardex contable, emitir reportes PDF y Excel, auditar existencias físicas y sincronizar datos en tiempo real con Firebase Realtime Database bajo un paradigma offline-first.

La solución combina Room Database como fuente de verdad local, Firebase como capa de nube multiusuario, Jetpack Compose con Material Design 3 para la interfaz, y una arquitectura MVVM con StateFlow y corrutinas. El sistema implementa roles empresariales (Administrador, Auditor y Visor), permisos granulares, notificaciones con sonido y vibración, temas visuales dinámicos y exportación profesional de documentos.

Este documento está dirigido a evaluadores académicos, equipos técnicos y stakeholders empresariales que requieran comprender tanto la arquitectura como el funcionamiento operativo del Inventario Agrícola.
""")

    # 2 INTRODUCCIÓN
    add_heading(doc, "2. Introducción", 1)
    add_heading(doc, "2.1 Contexto del Problema", 2)
    add_para(doc, """
Las empresas agrícolas gestionan inventarios complejos: fertilizantes, herramientas, repuestos, insumos y materiales distribuidos en fincas y bodegas. Los registros en papel o hojas de cálculo aisladas generan duplicidad, desfase de stock, imposibilidad de auditar y pérdida de información cuando no hay conectividad en campo.

El Inventario Agrícola responde a esta necesidad centralizando el control multibodega en un dispositivo Android, con operación offline y sincronización automática al recuperar red.
""")
    add_heading(doc, "2.2 Objetivos del Sistema", 2)
    add_table(doc, ["Objetivo", "Descripción"], [
        ["Control multibodega", "Administrar productos por bodega con códigos únicos"],
        ["Trazabilidad", "Kardex automático por cada movimiento"],
        ["Offline-first", "Operar sin internet y sincronizar después"],
        ["Seguridad", "Roles, permisos y auditoría de acciones"],
        ["Reportería", "Exportación PDF y Excel empresarial"],
        ["Alertas", "Stock bajo con notificaciones sonoras"],
    ])

    # 3 TECNOLOGÍAS
    add_heading(doc, "3. Stack Tecnológico", 1)
    techs = [
        ("Kotlin", "Lenguaje oficial Android. Tipado seguro, corrutinas nativas, interoperabilidad Java.", "Todo el código fuente del sistema."),
        ("Android Studio", "IDE oficial. Emuladores, profiler, Compose Preview, Gradle.", "Entorno de desarrollo y compilación."),
        ("Jetpack Compose", "UI declarativa moderna. Recomposición eficiente, menos XML.", "Todas las pantallas: login, dashboard, inventario, etc."),
        ("Material Design 3", "Sistema de diseño Google. ColorScheme, tipografía, componentes.", "Temas verde/azul/morado/naranja/oscuro."),
        ("MVVM", "Separación View-ViewModel-Model. Testabilidad y mantenimiento.", "Cada módulo tiene ViewModel dedicado."),
        ("StateFlow / MutableStateFlow", "Flujos reactivos de estado UI.", "SessionManager, AppThemeState, listas en ViewModels."),
        ("Coroutines", "Concurrencia ligera sin bloquear UI.", "Room, Firebase, export PDF en Dispatchers.IO."),
        ("Room Database v30", "ORM SQLite con DAOs, Flow, migraciones.", "17 entidades, inventario_db."),
        ("Firebase RTDB", "Base NoSQL en tiempo real, listeners.", "Sincronización multibodega en la nube."),
        ("Repository Pattern", "Abstracción de fuentes de datos.", "FirebaseRepository, InventoryRepository."),
        ("Navigation Compose", "Grafo de navegación tipado.", "NavGraph.kt con 40+ rutas."),
        ("Apache POI 5.2.5", "Generación Excel .xlsx.", "Export inventario, entradas, salidas, kardex."),
        ("PdfDocument / Canvas", "PDF nativo Android.", "Reportes con BrandingExports corporativo."),
    ]
    add_table(doc, ["Tecnología", "Descripción", "Uso en el sistema"], techs)

    # 4 ARQUITECTURA
    add_heading(doc, "4. Arquitectura del Sistema", 1)
    add_para(doc, """
La arquitectura sigue el patrón MVVM en capas. La UI (Compose) observa StateFlow del ViewModel. El ViewModel delega operaciones de negocio a servicios y repositorios. Room persiste localmente de forma inmediata; Firebase recibe push cuando hay conectividad. SyncQueueManager reintenta operaciones fallidas.
""")
    add_image(doc, ASSETS / "01_arquitectura_mvvm.png")
    add_image(doc, ASSETS / "02_offline_first.png")
    add_para(doc, """
Flujo de datos: (1) Usuario interactúa con pantalla Compose. (2) ViewModel procesa en viewModelScope. (3) MovimientoInventarioService escribe Room. (4) Si online → FirebaseRepository.guardar*. (5) Si offline → SyncQueueManager.enqueue. (6) OfflineSyncObserver detecta red → syncAllPendingAndBidirectional().
""")

    # 5 MÓDULOS
    add_heading(doc, "5. Módulos Funcionales", 1)
    modules = [
        ("Login", "Autenticación Room/Firebase, SessionManager, recuperación OTP"),
        ("Dashboard", "KPIs, gráficas Vico, alertas, predicción, presupuesto"),
        ("Inventario", "Consulta productos, filtros, export — sin edición directa"),
        ("Entradas", "Abastecimiento: producto nuevo/existente, factura, kardex"),
        ("Salidas", "Consumo operativo, vales, validación stock, alertas"),
        ("Kardex", "Historial movimientos: ENTRADA, SALIDA, TRASLADO, AJUSTE"),
        ("Facturas", "CRUD facturas con detalle, totales, export"),
        ("Vales", "Vales multilínea con DetalleVale e integración salidas"),
        ("Categorías", "Prefijos automáticos, correlativo atómico, sync Firebase"),
        ("Auditoría", "Conteo físico, ajustes automáticos, gráficas"),
        ("Stock bajo", "Productos bajo mínimo, export, alertas sonoras"),
        ("Configuración", "Perfil, temas, notificaciones, backups, papelera"),
        ("Usuarios/Roles", "ADMIN, AUDITOR, VISOR — RoleManager + permisos"),
        ("Reportes", "PDF y Excel por módulo con branding corporativo"),
    ]
    add_table(doc, ["Módulo", "Funcionalidad principal"], modules)

    # 6 FUNCIONAMIENTO
    add_heading(doc, "6. Funcionamiento Operativo", 1)
    add_heading(doc, "6.1 Registro de Entrada", 2)
    add_para(doc, """
El usuario accede a CrearEntradasScreen. Selecciona producto existente o crea uno nuevo con categoría (genera código TOR-0001). Completa cantidad, costo, proveedor y factura. EntradaViewModel invoca MovimientoInventarioService.registrarEntrada(): incrementa stock, calcula costo promedio ponderado, inserta Entrada, crea Factura+DetalleFactura, genera Kardex KAR-ENT-{id} y sincroniza Firebase. NotificationHelper emite alerta sonora en canal entradas_channel.
""")
    add_heading(doc, "6.2 Registro de Salida", 2)
    add_image(doc, ASSETS / "04_flujo_entradas_salidas.png")
    add_para(doc, """
SalidaViewModel.registrarSalidaCompleta() valida stock suficiente, decrementa cantidad, actualiza status (ACTIVO/STOCK_BAJO/SIN_STOCK), inserta Salida y Kardex KAR-SAL-{id}. Si cantidad <= stockMinimo, dispara notificación STOCK_BAJO con sonido de alerta en stock_channel.
""")
    add_heading(doc, "6.3 Kardex y Stock", 2)
    add_para(doc, """
Cada movimiento genera registro Kardex con saldoAnterior, cantidad, saldoNuevo, tipoMovimiento y referencia (factura/vale). El stock del producto es la fuente de verdad actualizada atómicamente en Room antes del push Firebase.
""")

    # 7 CÓDIGOS AUTOMÁTICOS
    add_heading(doc, "7. Sistema de Códigos Automáticos", 1)
    add_image(doc, ASSETS / "07_codigos_automaticos.png")
    add_para(doc, """
Formato: PREFIJO-CORRELATIVO (ej. TOR-0001, ACE-0002). CategoriaViewModel.generarCodigoProducto() ejecuta incrementarCorrelativo() atómico en SQL, sincroniza categoría a Firebase y construye código con CodigoGenerator. Regex validación: ^[A-Za-z]{2,4}-\\d{4}$. Previene duplicados por bodega+codigo único en ProductoDao.
""")

    # 8 BASE DE DATOS
    add_heading(doc, "8. Base de Datos Room", 1)
    add_image(doc, ASSETS / "03_erd_database.png")
    add_table(doc, ["Entidad", "Tabla", "Clave principal"], [
        ("Usuario", "usuarios", "id (auto)"),
        ("Bodega", "bodegas", "id (UUID)"),
        ("Producto", "productos", "id + codigo+bodegaId"),
        ("Categoria", "categorias", "id + prefijo"),
        ("Entrada/Salida", "entradas/salidas", "id (auto)"),
        ("Kardex", "kardex", "id (auto)"),
        ("Factura", "facturas", "id + numeroFactura"),
        ("Auditoria", "auditorias", "id (auto)"),
        ("Vale", "vale", "idVale (auto)"),
    ])
    add_para(doc, """
Versión 30 de appdatabase.kt. Soft-delete con isDeleted/deletionDate en entidades principales. 17 DAOs con Flow reactivo para UI. Relación lógica: Bodega 1:N Producto, Producto 1:N Entrada/Salida/Kardex, Categoria genera prefijos para Producto.
""")

    # 9 FIREBASE
    add_heading(doc, "9. Firebase Realtime Database", 1)
    add_image(doc, ASSETS / "05_firebase_sync.png")
    add_heading(doc, "9.1 Problemas Detectados y Soluciones", 2)
    add_table(doc, ["Problema", "Causa", "Solución implementada"], [
        ("Firebase no sincronizaba", "Push manual incompleto", "CloudSyncManager bidireccional + queue"),
        ("Categorías no aparecían", "Listeners incompletos", "ValueEventListener en CategoriaViewModel"),
        ("Room guardaba, Firebase no", "Sin cola offline", "SyncQueueManager con 5 reintentos"),
        ("Datos duplicados", "Conflictos merge", "ConflictResolver last-write-wins"),
        ("Realtime incompleto", "Sin observer red", "OfflineSyncObserver + NetworkCallback"),
    ])

    # 10 MEJORAS
    add_heading(doc, "10. Mejoras y Correcciones Realizadas", 1)
    add_para(doc, """
Separación Entradas (abastecimiento) vs Salidas (consumo). Inventario solo consulta. Formularios con scroll y autocompletado. Dashboard rediseñado con Material 3 y gráficas Vico visibles. Temas dinámicos (verde, azul, morado, naranja, oscuro) con glassmorphism. Notificaciones con 4 canales IMPORTANCE_HIGH y sonidos diferenciados. Roles simplificados: ADMIN, AUDITOR, VISOR. Papelera de recuperación. Auditoría estable con resolución codigoBodega.
""")

    # 11 INTERFAZ
    add_heading(doc, "11. Interfaz de Usuario", 1)
    for img, cap in [
        ("ui_login.png", "Figura 1. Pantalla de Login"),
        ("ui_dashboard.png", "Figura 2. Dashboard ejecutivo con KPIs"),
        ("ui_inventario.png", "Figura 3. Módulo Inventario (consulta)"),
        ("ui_entradas.png", "Figura 4. Módulo Entradas / Abastecimiento"),
        ("ui_salidas.png", "Figura 5. Módulo Salidas / Consumo"),
    ]:
        add_image(doc, ASSETS / img, Inches(2.8))
        cp = doc.add_paragraph(cap); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True

    # 12 DIAGRAMAS ADICIONALES
    add_heading(doc, "12. Diagramas Complementarios", 1)
    for img in ["06_roles_permisos.png", "08_navgraph.png", "09_stock_bajo.png", "10_exportacion.png"]:
        add_image(doc, ASSETS / img)

    # 13 CÓDIGO
    add_heading(doc, "13. Código Fuente Explicado", 1)
    add_heading(doc, "13.1 ViewModel con StateFlow", 2)
    code1 = doc.add_paragraph()
    code1.add_run("""
// EntradaViewModel.kt — patrón MVVM
fun agregarEntrada(entrada: Entrada, productoNuevo: Producto? = null) {
    viewModelScope.launch(Dispatchers.IO) {
        when (val r = registrarEntradaCompleta(entrada, productoNuevo)) {
            is ResultadoMovimiento.EntradaOk -> NotificationHelper.registrar(...)
        }
    }
}
""").font.name = "Courier New"; code1.runs[0].font.size = Pt(9)
    add_para(doc, """
El ViewModel encapsula la lógica fuera de la UI. viewModelScope cancela corrutinas al destruir pantalla. Dispatchers.IO evita bloquear hilo principal en operaciones de base de datos.
""")
    add_heading(doc, "13.2 DAO Room", 2)
    code2 = doc.add_paragraph()
    code2.add_run("""
@Dao interface ProductoDao {
    @Query("SELECT * FROM productos WHERE bodegaId = :bodegaId")
    fun obtenerProductos(bodegaId: String): Flow<List<Producto>>
    @Insert(onConflict = OnConflictStrategy.REPLACE)
    suspend fun insertar(producto: Producto): Long
}
""").font.name = "Courier New"; code2.runs[0].font.size = Pt(9)
    add_heading(doc, "13.3 Sync Queue", 2)
    add_para(doc, """
SyncQueueManager persiste operaciones pendientes en SharedPreferences como JSON. OfflineManager.enqueueSync() verifica NET_CAPABILITY_VALIDATED. Al recuperar conexión, processPending() re-lee entidad de Room y ejecuta FirebaseRepository.guardar*().
""")

    # 14 PROBLEMAS
    add_heading(doc, "14. Problemas Resueltos", 1)
    add_table(doc, ["Problema", "Corrección"], [
        ("Entradas/salidas mezcladas", "Flujos separados + MovimientoInventarioService"),
        ("Dashboard oscuro ilegible", "MaterialTheme.colorScheme + temas dinámicos"),
        ("Gráficas invisibles", "Vico charts + colores theme-aware"),
        ("Formularios sin scroll", "LazyColumn/verticalScroll en Compose"),
        ("Stock bajo sin sonido", "NotificationChannelManager IMPORTANCE_HIGH + raw sounds"),
        ("Firebase desincronizado", "Bidirectional sync + listeners + queue"),
    ])

    # 15 CONCLUSIÓN
    add_heading(doc, "15. Conclusión", 1)
    add_para(doc, """
El Inventario Agrícola Multibodega constituye una solución empresarial completa para la gestión de inventarios en el sector agrícola. Integra tecnologías modernas de Android (Compose, Room, Firebase) bajo arquitectura MVVM offline-first, garantizando operación en campo sin conectividad y sincronización automática en oficina.

Beneficios: trazabilidad kardex, códigos automáticos, roles de seguridad, reportería PDF/Excel, alertas proactivas de stock y auditoría física con ajustes. Escalabilidad: estructura multibodega en Firebase, módulos desacoplados, cola de sincronización.

Mejoras futuras sugeridas: migraciones Room no destructivas, hash contraseñas en producción (PasswordHasher PLAINTEXT_DEV_MODE → SHA-256), pruebas instrumentadas, panel web administrativo, integración código de barras/QR y analítica predictiva avanzada con ML on-device.

La importancia del sistema radica en digitalizar procesos críticos del agro guatemalteco, reduciendo pérdidas por descontrol de inventario y elevando la toma de decisiones con datos en tiempo real.
""")
    add_heading(doc, "Referencias", 1)
    refs = [
        "Google. (2024). Jetpack Compose documentation. https://developer.android.com/jetpack/compose",
        "Google. (2024). Room persistence library. https://developer.android.com/training/data-storage/room",
        "Google. (2024). Firebase Realtime Database. https://firebase.google.com/docs/database",
        "JetBrains. (2024). Kotlin coroutines guide. https://kotlinlang.org/docs/coroutines-guide.html",
        "Apache Software Foundation. (2024). Apache POI. https://poi.apache.org/",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{ref}", style="Normal")

    out = OUT / "Informe_Tecnico_Inventario_Agro.docx"
    doc.save(out)
    return out


# ─── PDF generation ───────────────────────────────────────────────────────────

def build_pdf():
    out = OUT / "Informe_Tecnico_Inventario_Agro.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER,
                              fontName="Times-Roman", fontSize=22, textColor=colors.HexColor("#1B5E20"),
                              spaceAfter=20))
    styles.add(ParagraphStyle(name="BodyJustify", parent=styles["Normal"], alignment=TA_JUSTIFY,
                              fontName="Times-Roman", fontSize=12, leading=18, spaceAfter=10))
    styles.add(ParagraphStyle(name="H1", parent=styles["Heading1"], fontName="Times-Bold",
                              fontSize=16, textColor=colors.HexColor("#1B5E20"), spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2", parent=styles["Heading2"], fontName="Times-Bold",
                              fontSize=14, textColor=colors.HexColor("#2E7D32"), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="Caption", parent=styles["Normal"], alignment=TA_CENTER,
                              fontName="Times-Italic", fontSize=10, textColor=colors.grey))

    story = []

    # Portada
    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph(PROJECT, styles["CenterTitle"]))
    story.append(Paragraph(SUBTITLE, ParagraphStyle(name="Sub", alignment=TA_CENTER, fontName="Times-Roman", fontSize=14)))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(f"Documentación Técnica Integral<br/><br/>Autor: {AUTHOR}<br/>{INSTITUTION}<br/>{DATE}<br/>Versión 1.0",
                           ParagraphStyle(name="Meta", alignment=TA_CENTER, fontName="Times-Roman", fontSize=12, leading=18)))
    story.append(PageBreak())

    sections = [
        ("1. Resumen Ejecutivo", """
        El Inventario Agrícola Multibodega es una aplicación Android empresarial desarrollada en Kotlin que centraliza 
        el control de inventarios agrícolas en múltiples bodegas. Combina Room Database (offline-first), Firebase Realtime 
        Database (sincronización en la nube), Jetpack Compose con Material 3, arquitectura MVVM, StateFlow, corrutinas, 
        generación de kardex, reportes PDF/Excel, auditoría física, roles de seguridad y notificaciones sonoras profesionales.
        """),
        ("2. Introducción", """
        Las empresas agrícolas requieren control riguroso de insumos distribuidos geográficamente. Este sistema digitaliza
        entradas, salidas, stock, facturas, vales y traslados entre bodegas, operando sin conexión y sincronizando 
        automáticamente al detectar red disponible mediante OfflineSyncObserver y SyncQueueManager.
        """),
        ("3. Stack Tecnológico", """
        Kotlin 2.1.0 · Android SDK 35 · Jetpack Compose BOM 2024.11 · Material 3 · Room 2.6.1 · Firebase BOM 33.7.0 · 
        Navigation Compose 2.8.4 · Apache POI 5.2.5 · Vico Charts 1.13.1 · Coroutines Play Services 1.9.0 · Coil 2.6.0.
        MVVM separa UI (Compose) de lógica (ViewModel) y datos (Repository/DAO). StateFlow expone estado reactivo a la UI.
        """),
        ("4. Arquitectura MVVM y Offline-First", """
        Capas: UI Compose → ViewModel → MovimientoInventarioService / Repository → Room ↔ Firebase. 
        Escritura inmediata en SQLite local garantiza respuesta instantánea. Push a Firebase cuando hay red; 
        cola de reintentos cuando no. CloudSyncManager ejecuta sincronización bidireccional completa.
        """),
        ("5. Módulos del Sistema", """
        Login · Dashboard KPIs · Inventario (consulta) · Entradas (abastecimiento) · Salidas (consumo) · Kardex · 
        Facturas · Vales · Categorías · Auditoría · Stock bajo · Presupuesto · Reportes operativos · Configuración · 
        Usuarios · Papelera · Logs · Temas visuales · Notificaciones.
        """),
        ("6. Funcionamiento: Entradas y Salidas", """
        ENTRADA: CrearEntradasScreen → EntradaViewModel → registrarEntrada() → stock↑, kardex, factura, Firebase, notificación.
        SALIDA: CrearSalidasScreen → SalidaViewModel → registrarSalida() → stock↓, kardex, validación, alerta stock bajo si aplica.
        """),
        ("7. Sistema de Códigos Automáticos", """
        Formato PREFIJO-#### (TOR-0001). Categoria almacena prefijo y correlativoActual. incrementarCorrelativo() atómico en SQL.
        CodigoGenerator construye prefijo desde nombre categoría. Sincronización Firebase de categoría tras incremento.
        """),
        ("8. Base de Datos Room v30", """
        17 entidades: usuarios, bodegas, productos, categorias, entradas, salidas, facturas, detalle_factura, kardex, 
        vales, detalle_vale, traslados, auditorias, logs, app_notificaciones. Soft-delete. DAOs con Flow<List>.
        """),
        ("9. Firebase y Sincronización", """
        Estructura: /bodegas/{codigo}/{bodegaId}/productos|entradas|salidas|kardex|facturas|vales|auditorias.
        Problemas resueltos: listeners ValueEventListener, push automático post-Room, pull CloudSyncManager, 
        ConflictResolver merge, persistencia Firebase habilitada.
        """),
        ("10. Mejoras Implementadas", """
        Separación entradas/salidas · Inventario solo lectura · Temas Material 3 dinámicos · Glassmorphism · 
        Dashboard gráficas Vico · Notificaciones 4 canales con sonido · Roles ADMIN/AUDITOR/VISOR · 
        Botón regresar en pantallas · Papelera recuperación · Export PDF async.
        """),
        ("11. Seguridad y Roles", """
        RoleManager matriz permisos. SessionManager StateFlow sesión. ModuleRouteGuard en NavGraph. 
        ADMIN: operación completa. AUDITOR: supervisión total. VISOR: solo lectura.
        """),
        ("12. Exportación y Reportería", """
        Apache POI genera Excel .xlsx. PdfDocument + Canvas genera PDF nativo. BrandingExports aplica logo y paleta verde.
        FileProvider comparte archivos. Permiso EXPORTAR requerido.
        """),
        ("13. Notificaciones", """
        NotificationChannelManager: entradas_channel, salidas_channel, stock_channel, auditoria_channel.
        IMPORTANCE_HIGH, sonidos raw diferenciados, vibración configurable en perfil. POST_NOTIFICATIONS Android 13+.
        """),
        ("14. Código Fuente — Patrones Clave", """
        ViewModel + viewModelScope + Dispatchers.IO. DAO suspend/Flow. FirebaseRepository guardar* con await().
        SyncQueueManager JSON queue SharedPreferences. NotificationHelper registrar → Room + Firebase + push local.
        """),
        ("15. Conclusión", """
        El sistema Inventario Agrícola demostrado es escalable, profesional y apto para entornos empresariales reales. 
        Combina robustez offline-first con sincronización cloud, trazabilidad kardex, seguridad por roles y 
        experiencia de usuario moderna. Mejoras futuras: migraciones Room, hash producción, tests E2E, panel web, QR/ML.
        """),
    ]

    diagram_map = {
        "4. Arquitectura MVVM y Offline-First": ["01_arquitectura_mvvm.png", "02_offline_first.png"],
        "6. Funcionamiento: Entradas y Salidas": ["04_flujo_entradas_salidas.png"],
        "7. Sistema de Códigos Automáticos": ["07_codigos_automaticos.png"],
        "8. Base de Datos Room v30": ["03_erd_database.png"],
        "9. Firebase y Sincronización": ["05_firebase_sync.png"],
        "11. Seguridad y Roles": ["06_roles_permisos.png"],
        "12. Exportación y Reportería": ["10_exportacion.png"],
        "13. Notificaciones": ["09_stock_bajo.png"],
    }

    ui_imgs = ["ui_login.png", "ui_dashboard.png", "ui_inventario.png"]

    for title, body in sections:
        story.append(Paragraph(title, styles["H1"]))
        story.append(Paragraph(body.strip().replace("\n", " "), styles["BodyJustify"]))
        if title in diagram_map:
            for d in diagram_map[title]:
                p = ASSETS / d
                if p.exists():
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Image(str(p), width=5.5*inch, height=3*inch))
                    story.append(Spacer(1, 0.15*inch))
        if title.startswith("5."):
            for u in ui_imgs:
                p = ASSETS / u
                if p.exists():
                    story.append(Image(str(p), width=2.2*inch, height=3.2*inch))
            story.append(Spacer(1, 0.2*inch))
        story.append(Spacer(1, 0.1*inch))

    story.append(PageBreak())
    story.append(Paragraph("Referencias (APA)", styles["H1"]))
    for ref in [
        "Google. (2024). Jetpack Compose. developer.android.com/jetpack/compose",
        "Google. (2024). Room persistence library. developer.android.com/training/data-storage/room",
        "Google. (2024). Firebase Realtime Database. firebase.google.com/docs/database",
        "JetBrains. (2024). Kotlin coroutines. kotlinlang.org/docs/coroutines-guide.html",
    ]:
        story.append(Paragraph(ref, styles["BodyJustify"]))

    doc.build(story)
    return out


def main():
    os.environ["MPLCONFIGDIR"] = str(BASE / ".mpl_cache")
    print("Generando diagramas...")
    generate_all_diagrams()
    print("Generando Word...")
    docx_path = build_docx()
    print(f"  → {docx_path}")
    print("Generando PDF...")
    pdf_path = build_pdf()
    print(f"  → {pdf_path}")
    print("¡Informe técnico generado exitosamente!")


if __name__ == "__main__":
    main()
